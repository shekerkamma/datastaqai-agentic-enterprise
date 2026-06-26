#!/usr/bin/env python3
"""Build the Word report from the chapter spec files (headlines + speaker notes =
the written narrative) + a figures appendix."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

H=Path("/mnt/d/New folder/Antigravity-test/antigravity-skills/wedge_dossiers/book")
A=Path("/mnt/d/New folder/Antigravity-test/antigravity-skills/wedge_dossiers/analysis")
NAVY=RGBColor(0x0A,0x16,0x28); TEAL=RGBColor(0x00,0x9B,0x82)
CH=[("Chapter1-Executive-Thesis.md","Part I — Executive Thesis"),
    ("Chapter2-AI-Transformation.md","Part II — The AI Transformation"),
    ("Chapter3-Business-Use-Cases.md","Part III — Business Use Cases"),
    ("Chapter4a-Competitive.md","Part IV — Competitive Analysis (1/2)"),
    ("Chapter4b-Competitive.md","Part IV — Competitive Analysis (2/2)"),
    ("Chapter5-Business-Model.md","Part V — Business Model"),
    ("Chapter6-Path-to-5-10M.md","Part VI — Path to $5–10M"),
    ("Chapter7-Investment-Thesis.md","Part VII — Investment Thesis")]

def parse(md):
    txt=(H/md).read_text(encoding="utf-8")
    out=[]
    for chunk in re.split(r"\n## Slide ",txt)[1:]:
        head=chunk.splitlines()[0]
        m=re.match(r"\d+\s+—\s+(.*)",head) or re.match(r"\d+\s+—\s+(.*)",head)
        headline=head.split("—",1)[1].strip() if "—" in head else head
        num=head.split("—")[0].strip()
        mn=re.search(r"\*\*Speaker notes\.\*\*\s*\n+(.*)$",chunk,re.S)
        body=mn.group(1).strip() if mn else ""
        body=re.sub(r"\n+"," ",body).strip()
        out.append((num,headline,body))
    return out

doc=Document()
st=doc.styles["Normal"]; st.font.name="Calibri"; st.font.size=Pt(11)

# title page
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run("THE AGENTIC ENTERPRISE"); r.bold=True; r.font.size=Pt(30); r.font.color.rgb=NAVY
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run("From software-assisted work to AI-orchestrated work"); r.font.size=Pt(16); r.font.color.rgb=TEAL
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
s2.add_run("A Strategy Brief by DataStaqAI · 2026 · 48-slide companion deck").italic=True
doc.add_paragraph()
lead=doc.add_paragraph()
lead.add_run("Thesis. ").bold=True
lead.add_run("Enterprise work is moving from software-assisted (per-seat SaaS) to AI-orchestrated (agentic). "
  "Value will accrue not to the commoditizing model nor the seat-defending incumbent, but to whoever delivers the "
  "orchestration — the agent wired into the customer's systems of record, governed and operated. Every category has "
  "split into incumbents and AI-natives, and both leave the mid-market done-for-you delivery layer unowned. DataStaqAI "
  "occupies that layer. All figures in this report are source-traceable to a competitive research base of 25 agentic "
  "use cases and 40+ named vendors; constructed financial figures are labeled illustrative.")
doc.add_page_break()

# chapters
for md,title in CH:
    doc.add_heading(title,level=1)
    for num,headline,body in parse(md):
        h=doc.add_heading(f"Slide {num} — {headline}",level=2)
        if body: doc.add_paragraph(body)

# figures appendix
doc.add_page_break(); doc.add_heading("Appendix — Figures",level=1)
figs=[(A/"chart1_posture_mix.png","Delivery posture across 25 use cases"),
      (A/"chart2_price_floor_gap.png","Mid-market price-floor gap"),
      (A/"chart3_tam_by_usecase.png","Sourced 2025 TAM by use case"),
      (A/"chart4_opportunity_map.png","Opportunity map (density × pattern)"),
      (A/"chart5_funding_by_cluster.png","Disclosed funding by cluster"),
      (A/"chart6_valuation_leaders.png","Valuation leaders"),
      (A/"chart7_valuation_vs_arr.png","Valuation vs ARR"),
      (H/"chart_cagr.png","Agentic sub-market CAGRs"),
      (H/"chart_coverage.png","Mid-market coverage heat map"),
      (H/"chart_disruption.png","Disruption matrix"),
      (H/"chart_arr_ramp.png","Illustrative ARR ramp to $5–10M")]
for path,cap in figs:
    if path.exists():
        doc.add_picture(str(path),width=Inches(6.2))
        c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        rr=c.add_run("Figure: "+cap); rr.italic=True; rr.font.size=Pt(9)
        doc.add_paragraph()

out=H/"DataStaqAI-The-Agentic-Enterprise-Report.docx"
doc.save(str(out))
print("WORD:",out,"| paragraphs:",len(doc.paragraphs))
