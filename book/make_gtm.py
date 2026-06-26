#!/usr/bin/env python3
"""#3 — 90-day GTM execution plan (.docx + .md)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
H=Path("/mnt/d/New folder/Antigravity-test/antigravity-skills/wedge_dossiers/book")
NAVY=RGBColor(0x0A,0x16,0x28); TEAL=RGBColor(0x00,0x9B,0x82)
doc=Document(); doc.styles["Normal"].font.name="Calibri"; doc.styles["Normal"].font.size=Pt(11)
def H1(t): doc.add_heading(t,1)
def H2(t): doc.add_heading(t,2)
def P(t,bold=False,it=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=it; return p
def bullets(items):
    for i in items: doc.add_paragraph(i,style="List Bullet")
def table(headers,rows):
    t=doc.add_table(rows=1,cols=len(headers)); t.style="Light Grid Accent 1"
    for j,h in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=h
        for rn in c.paragraphs[0].runs: rn.bold=True
    for row in rows:
        cells=t.add_row().cells
        for j,v in enumerate(row): cells[j].text=str(v)
    return t

# title
tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("DataStaqAI — 90-Day GTM Execution Plan"); r.bold=True; r.font.size=Pt(24); r.font.color.rgb=NAVY
sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
sp.add_run("Make the strategy executable: vertical focus · trigger-event outbound · first deals · messaging").italic=True
doc.add_paragraph()

H1("0 · The 90-day goal")
bullets(["North star: prove the delivery motion and earn referenceable proof — not maximize logos.",
 "90-day targets (planning case): 1 BUILD win signed · 3–5 paid pilots live · 3 reference case studies · qualified pipeline = 4–6x quarterly capacity.",
 "Operating principle: lead with a BUILD wedge to win trust, then expand the landed account into recurring INTEGRATE retainers."])

H1("1 · Vertical focus (concentrate to compound)")
P("Primary vertical: Mid-market Manufacturing & Distribution (200–2,000 employees).",bold=True)
bullets(["Why: the densest stack of acute, fundable pain in one buyer — AP/invoice (IDP), CPQ/quoting, predictive ops, and low-code/legacy apps all live here.",
 "Compounding: a referenceable win + reusable library (AP agent, ERP connectors, CPQ agent) transfers directly to the next account in the same vertical."])
P("Secondary vertical: Mid-size Legal firms (45–250 attorneys).",bold=True)
bullets(["Why: sharpest price dislocation ('$20 tool vs $270 seat'), clear ICP and champion (Managing Partner / Legal Ops), referral-dense community, hallucination-governance is a paid differentiator.",
 "Sequencing: start outbound in Manufacturing in Weeks 1–6; open Legal in Weeks 6–12 once the first proof exists."])

H1("2 · The offer to lead with")
bullets(["LEAD (BUILD wedge): 'Low-code Escape' or 'Legacy Modernization' — owned code the client keeps; fixed-scope, weeks not quarters. Door-opener + trust.",
 "LAND (fast INTEGRATE pilot): one wired agent on the system of record (AP-into-ERP, CPQ, or SOC investigation) at $2–3k/mo, live in 2–4 weeks.",
 "EXPAND: convert the pilot account to a 2–3 use-case managed program ($5–8k/mo) within a quarter."])

H1("3 · Trigger-event outbound playbook")
P("Hunt forced buyers, not cold ones. Each row is a signal → where to find it → the hook → the entry offer.")
table(["Trigger signal","Where to source it","Message hook","Entry offer"],
 [["Salesforce CPQ End-of-Sale","RevOps on LinkedIn; CPQ communities; news","'Avoid the $200/seat Revenue Cloud rebuild'","Conversational quote agent pilot"],
  ["SOC-analyst / Tier-1 job posts","Job boards (LinkedIn, Indeed, Bayt)","'You're hiring to triage alerts an agent can investigate'","DEPLOY investigation agent on your SIEM"],
  ["Low-code license re-pricing","OutSystems/Mendix forums; LinkedIn gripes","'Your low-code bill now exceeds your eng payroll'","BUILD: low-code escape to owned code"],
  ["Failed OCR / manual AP","r/Accounting; AP-manager posts; RFPs","'Your OCR choked on half your invoices'","No-template AP agent into your ERP"],
  ["Legal 'which-AI-tools' review","Law.com; legal-ops communities; firm news","'Priced out of Harvey, exposed on hallucinations'","Citator-verified research/draft pilot"]])

H1("4 · First 5 target-account profiles (ICP archetypes)")
P("Archetypes to source against (not named targets) — qualify in only when the trigger is present.")
table(["#","Profile","Trigger","Champion","Entry offer","Expected first ACV"],
 [["1","~500-person industrial manufacturer, Dynamics/SAP ERP, 10k+ invoices/yr","Failed OCR; AP can't keep up","Controller / VP Finance","AP agent + 3-way match","$30–40k/yr"],
  ["2","Mid-market distributor on legacy Salesforce CPQ","CPQ End-of-Sale","RevOps / VP Sales Ops","Conversational quote agent","$30–40k/yr"],
  ["3","Manufacturer with a 3+ yr OutSystems/Mendix app + rising bill","Low-code re-pricing","CTO / VP Engineering","BUILD: low-code escape","$60–150k fixed"],
  ["4","100-attorney commercial firm evaluating legal AI","'Which tools' review","Managing Partner / Legal Ops","Verified research/draft pilot","$30–48k/yr"],
  ["5","Mid-market SaaS/fintech SOC, 4–8 analysts, hiring L1","SOC-analyst job post","CISO / SecOps lead","DEPLOY investigation agent","$36–60k/yr"]])

H1("5 · Messaging by persona")
table(["Persona","Their pain (1 line)","Our hook","Proof to cite"],
 [["CTO / VP Engineering","Trapped in low-code; can't add AI; rewrite risk","'Out of the cage into code you own, in a quarter'","70% of manual rewrites fail; Mechanical Orchard = Global-2000-only"],
  ["CFO / Controller","AP is 3–5 FTE typing; close slips","'AP stops being a hiring problem'","8–12 min/invoice; Ramp/Maximor outcomes"],
  ["CISO / SecOps","Hiring to triage 90%-FP alerts","'Investigate 100%, often self-funded via SIEM cuts'","r/cybersecurity 'drowning in false positives'"],
  ["Managing Partner (Legal)","Priced out of Harvey; sanction risk","'Punch above your size, verified, governed'","'$20 beats $270'; firms sanctioned for fake cites"],
  ["RevOps / VP Sales Ops","Reps quote in Excel; CPQ EOS","'Quotes at conversation speed, off the dying CPQ'","CPQ End-of-Sale strands 6,000+ orgs"]])

H1("6 · 13-week timeline")
table(["Phase","Weeks","Focus","Exit milestone"],
 [["Set up","1–2","ICP list build, trigger-monitoring, 2 offers packaged, outbound assets, 1 reusable BUILD scaffold","20–30 sourced accounts; outbound live"],
  ["Land","3–6","Trigger outbound (Manufacturing); book pilots; sign 1 BUILD","2–3 pilots live + 1 BUILD signed"],
  ["Deliver + open Legal","7–10","Deliver pilots; start the BUILD; open Legal outbound","First outcomes shipped; Legal pipeline started"],
  ["Prove + expand","11–13","Case studies; expand 1 pilot to multi-function; ask for referrals","3 references; 1 expansion; pipeline = 4–6x capacity"]])

H1("7 · Metrics & founder time")
bullets(["Leading: # triggered accounts sourced/wk · outbound→meeting rate · meeting→pilot rate.",
 "Lagging: pilots live · BUILD signed · references · expansion ARR · CAC payback (<6–9 mo).",
 "Founder time (illustrative): ~50% selling/delivery on first deals, ~30% building the reusable library, ~20% recruiting the first pod.",
 "Capital-efficient by design: founder-led + referral motion keeps cash CAC low; hire the first pod only behind signed recurring ARR."])

H1("8 · Risks to manage in the first 90 days")
bullets(["Body-shop drift — productize the 2 offers from day one; bank every reusable asset.",
 "Over-broad targeting — stay in the 2 verticals; decline out-of-ICP deals.",
 "Delivery overcommit — cap concurrent pilots to protect quality (the first references are everything)."])

out=H/"DataStaqAI-90-Day-GTM-Plan.docx"; doc.save(str(out)); print("GTM:",out)
